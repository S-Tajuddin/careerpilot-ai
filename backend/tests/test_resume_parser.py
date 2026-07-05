"""
CareerPilot AI — Resume Parser Tests
Tests for resume text extraction, LLM parsing, profile update,
and resume-based search query generation.
"""

import json
import os
import tempfile
from pathlib import Path
from unittest.mock import MagicMock, patch, AsyncMock

import pytest

# ─── Text Extraction Tests ────────────────────────────────────────────────


class TestTextExtraction:
    """Test file text extraction from PDF, DOCX, TXT."""

    def test_extract_txt(self):
        """Test extracting text from a plain text file."""
        from app.services.resume_parser import ResumeParserService

        parser = ResumeParserService()
        with tempfile.NamedTemporaryFile(mode='w', suffix='.txt', delete=False, encoding='utf-8') as f:
            f.write("John Doe\nAEM Developer\n8 years experience\nSkills: AEM, Java, Sling")
            f.flush()
            text = parser.extract_text_from_file(f.name)
        os.unlink(f.name)

        assert "John Doe" in text
        assert "AEM Developer" in text
        assert "8 years experience" in text

    def test_extract_docx(self):
        """Test extracting text from a DOCX file."""
        from app.services.resume_parser import ResumeParserService
        from docx import Document

        parser = ResumeParserService()

        with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as f:
            doc = Document()
            doc.add_paragraph("Jane Smith")
            doc.add_paragraph("Senior AEM Architect")
            doc.add_paragraph("Skills: AEM, EDS, Java, Sling, OSGi")
            doc.save(f.name)
            f.flush()

            text = parser.extract_text_from_file(f.name)

        os.unlink(f.name)
        assert "Jane Smith" in text
        assert "Senior AEM Architect" in text
        assert "AEM" in text

    def test_unsupported_format(self):
        """Test that unsupported formats raise ValueError."""
        from app.services.resume_parser import ResumeParserService

        parser = ResumeParserService()
        with pytest.raises(ValueError, match="Unsupported file format"):
            parser.extract_text_from_file("resume.xyz")

    def test_empty_file_returns_empty(self):
        """Test that empty text files return empty string."""
        from app.services.resume_parser import ResumeParserService

        parser = ResumeParserService()
        with tempfile.NamedTemporaryFile(mode='w', suffix='.txt', delete=False, encoding='utf-8') as f:
            f.write("")
            f.flush()
            text = parser.extract_text_from_file(f.name)
        os.unlink(f.name)
        assert text == ""


# ─── LLM Response Parsing Tests ────────────────────────────────────────────


class TestLLMResponseParsing:
    """Test parsing of LLM responses into structured data."""

    def test_parse_valid_json(self):
        """Test parsing a valid JSON response."""
        from app.services.resume_parser import ResumeParserService

        parser = ResumeParserService()
        response = '{"full_name": "John Doe", "skills": ["AEM", "Java"], "experience_years": 8}'
        result = parser._parse_llm_response(response)

        assert result["full_name"] == "John Doe"
        assert "AEM" in result["skills"]
        assert result["experience_years"] == 8

    def test_parse_json_with_markdown_fences(self):
        """Test parsing JSON wrapped in markdown code fences."""
        from app.services.resume_parser import ResumeParserService

        parser = ResumeParserService()
        response = '```json\n{"full_name": "Jane", "skills": ["EDS"]}\n```'
        result = parser._parse_llm_response(response)

        assert result["full_name"] == "Jane"
        assert "EDS" in result["skills"]

    def test_parse_empty_response(self):
        """Test that empty response returns empty dict."""
        from app.services.resume_parser import ResumeParserService

        parser = ResumeParserService()
        result = parser._parse_llm_response("")
        assert result == {}

    def test_parse_invalid_json(self):
        """Test that invalid JSON returns empty dict."""
        from app.services.resume_parser import ResumeParserService

        parser = ResumeParserService()
        result = parser._parse_llm_response("This is not JSON at all")
        assert result == {}

    def test_parse_json_embedded_in_text(self):
        """Test parsing JSON embedded within text."""
        from app.services.resume_parser import ResumeParserService

        parser = ResumeParserService()
        response = 'Here is the parsed data:\n{"full_name": "Test User", "experience_years": 5}\nEnd of output.'
        result = parser._parse_llm_response(response)

        assert result["full_name"] == "Test User"
        assert result["experience_years"] == 5


# ─── Profile Update Tests ──────────────────────────────────────────────────


class TestProfileUpdate:
    """Test updating profile from parsed resume data."""

    def _make_profile(self):
        """Create a mock profile object."""
        profile = MagicMock()
        profile.full_name = None
        profile.email = None
        profile.phone = None
        profile.summary = None
        profile.skills_list = []
        profile.skills = "[]"
        profile.experience_years = 8.0
        profile.current_role = None
        profile.target_role = None
        profile.location = None
        profile.remote_preference = "any"
        profile.resume_text = None
        profile.resume_file_path = None
        profile.updated_at = None

        # Make skills_list work as property
        type(profile).skills_list = property(
            lambda self: json.loads(self.skills) if self.skills else [],
            lambda self, val: setattr(self, 'skills', json.dumps(val) if val else "[]")
        )

        return profile

    def test_update_name_when_empty(self):
        """Test that name is updated from resume when profile name is empty."""
        from app.services.resume_parser import ResumeParserService

        parser = ResumeParserService()
        profile = self._make_profile()
        parsed = {"full_name": "John Doe", "skills": ["AEM", "Java"], "experience_years": 10}

        result = parser.update_profile_from_parsed_resume(
            parsed=parsed,
            resume_text="Resume text here",
            file_path="/resumes/test.pdf",
            profile=profile,
            db=MagicMock(),
        )

        assert profile.full_name == "John Doe"
        assert "full_name" in result["updated_fields"]

    def test_email_phone_always_overwritten(self):
        """Test that email and phone from resume always overwrite profile."""
        from app.services.resume_parser import ResumeParserService

        parser = ResumeParserService()
        profile = self._make_profile()
        profile.email = "old@email.com"
        profile.phone = "1111111111"

        result = parser.update_profile_from_parsed_resume(
            parsed={"email": "new@email.com", "phone": "9999999999", "skills": ["AEM"]},
            resume_text="Resume",
            file_path="/resumes/test.pdf",
            profile=profile,
            db=MagicMock(),
        )

        assert profile.email == "new@email.com"
        assert profile.phone == "9999999999"

    def test_location_always_overwritten(self):
        """Test that location from resume always overwrites profile."""
        from app.services.resume_parser import ResumeParserService

        parser = ResumeParserService()
        profile = self._make_profile()
        profile.location = "Bangalore, India"

        result = parser.update_profile_from_parsed_resume(
            parsed={"location": "Hyderabad, India", "skills": ["AEM"]},
            resume_text="Resume",
            file_path="/resumes/test.pdf",
            profile=profile,
            db=MagicMock(),
        )

        assert profile.location == "Hyderabad, India"

    def test_skills_are_merged(self):
        """Test that resume skills are merged with existing profile skills."""
        from app.services.resume_parser import ResumeParserService

        parser = ResumeParserService()
        profile = self._make_profile()
        profile.skills = json.dumps(["AEM", "Docker"])  # Existing skills

        parsed = {"skills": ["Java", "Sling", "AEM"]}  # Resume has AEM + new skills

        result = parser.update_profile_from_parsed_resume(
            parsed=parsed,
            resume_text="Resume text",
            file_path="/resumes/test.pdf",
            profile=profile,
            db=MagicMock(),
        )

        # Skills should be merged
        updated_skills = json.loads(profile.skills) if isinstance(profile.skills, str) else profile.skills_list
        assert "AEM" in updated_skills
        assert "Java" in updated_skills
        assert "Sling" in updated_skills
        assert "Docker" in updated_skills

    def test_resume_text_and_path_always_updated(self):
        """Test that resume_text and resume_file_path are always set."""
        from app.services.resume_parser import ResumeParserService

        parser = ResumeParserService()
        profile = self._make_profile()

        result = parser.update_profile_from_parsed_resume(
            parsed={"skills": ["AEM"]},
            resume_text="My resume content",
            file_path="/resumes/resume.pdf",
            profile=profile,
            db=MagicMock(),
        )

        assert profile.resume_text == "My resume content"
        assert profile.resume_file_path == "/resumes/resume.pdf"
        assert "resume_text" in result["updated_fields"]

    def test_experience_updated(self):
        """Test that experience years are updated from resume."""
        from app.services.resume_parser import ResumeParserService

        parser = ResumeParserService()
        profile = self._make_profile()

        result = parser.update_profile_from_parsed_resume(
            parsed={"experience_years": 12, "skills": ["AEM"]},
            resume_text="Resume",
            file_path="/resumes/test.pdf",
            profile=profile,
            db=MagicMock(),
        )

        assert profile.experience_years == 12.0

    def test_target_role_from_target_roles(self):
        """Test that target_role is set from the first target_roles entry."""
        from app.services.resume_parser import ResumeParserService

        parser = ResumeParserService()
        profile = self._make_profile()

        result = parser.update_profile_from_parsed_resume(
            parsed={"target_roles": ["AEM Architect", "Senior AEM Developer"], "skills": ["AEM"]},
            resume_text="Resume",
            file_path="/resumes/test.pdf",
            profile=profile,
            db=MagicMock(),
        )

        assert profile.target_role == "AEM Architect"

    def test_name_overwritten_by_resume(self):
        """Test that existing profile name IS overwritten by resume (resume is source of truth)."""
        from app.services.resume_parser import ResumeParserService

        parser = ResumeParserService()
        profile = self._make_profile()
        profile.full_name = "Existing Name"

        result = parser.update_profile_from_parsed_resume(
            parsed={"full_name": "Resume Name", "skills": ["AEM"]},
            resume_text="Resume",
            file_path="/resumes/test.pdf",
            profile=profile,
            db=MagicMock(),
        )

        assert profile.full_name == "Resume Name"
        assert "full_name" in result["updated_fields"]


# ─── Search Query Generation Tests ─────────────────────────────────────────


class TestSearchQueryGeneration:
    """Test generating search queries from resume data."""

    def _make_profile(self, skills=None, target_role="AEM Architect", experience_years=10, location="Hyderabad, India"):
        profile = MagicMock()
        profile.skills_list = skills or ["AEM", "Adobe Experience Manager", "EDS", "Java", "Sling", "OSGi", "HTL"]
        profile.target_role = target_role
        profile.current_role = "Senior AEM Developer"
        profile.experience_years = experience_years
        profile.location = location
        profile.resume_text = "Some resume text"
        return profile

    def test_generates_queries_from_profile(self):
        """Test that search queries are generated from profile data."""
        from app.services.resume_parser import ResumeParserService

        parser = ResumeParserService()
        profile = self._make_profile()

        queries = parser.generate_search_queries_from_resume(profile)

        assert len(queries) > 0
        # Should include target role query
        assert any("AEM Architect" in q["query"] for q in queries)

    def test_includes_remote_search(self):
        """Test that remote search queries are included."""
        from app.services.resume_parser import ResumeParserService

        parser = ResumeParserService()
        profile = self._make_profile()

        queries = parser.generate_search_queries_from_resume(profile)

        assert any("remote" in q["query"].lower() for q in queries)

    def test_senior_level_queries_for_experienced(self):
        """Test that senior/architect queries are generated for 8+ years."""
        from app.services.resume_parser import ResumeParserService

        parser = ResumeParserService()
        profile = self._make_profile(experience_years=10)

        queries = parser.generate_search_queries_from_resume(profile)

        # Should have architect/lead queries
        assert any("architect" in q["query"].lower() or "lead" in q["query"].lower() for q in queries)

    def test_eds_queries_when_skills_include_eds(self):
        """Test that EDS-specific queries are generated when EDS is in skills."""
        from app.services.resume_parser import ResumeParserService

        parser = ResumeParserService()
        profile = self._make_profile(skills=["AEM", "EDS", "Edge Delivery Services", "Java"])

        queries = parser.generate_search_queries_from_resume(profile)

        assert any("eds" in q["query"].lower() or "edge delivery" in q["query"].lower() for q in queries)

    def test_no_duplicate_queries(self):
        """Test that duplicate queries are removed."""
        from app.services.resume_parser import ResumeParserService

        parser = ResumeParserService()
        profile = self._make_profile()

        queries = parser.generate_search_queries_from_resume(profile)
        query_texts = [q["query"].lower() for q in queries]

        assert len(query_texts) == len(set(query_texts)), "Duplicate queries found"

    def test_max_8_queries(self):
        """Test that no more than 8 queries are generated."""
        from app.services.resume_parser import ResumeParserService

        parser = ResumeParserService()
        profile = self._make_profile()

        queries = parser.generate_search_queries_from_resume(profile)

        assert len(queries) <= 8

    def test_location_based_query(self):
        """Test that location-specific queries are generated."""
        from app.services.resume_parser import ResumeParserService

        parser = ResumeParserService()
        profile = self._make_profile(location="Bangalore, India")

        queries = parser.generate_search_queries_from_resume(profile)

        assert any("Bangalore" in q["query"] for q in queries)


# ─── Resume Status Tests ───────────────────────────────────────────────────


class TestResumeStatus:
    """Test resume status reporting."""

    def test_no_resume_status(self):
        """Test status when no resume is uploaded."""
        from app.services.resume_parser import ResumeParserService

        parser = ResumeParserService()
        profile = MagicMock()
        profile.resume_text = None
        profile.resume_file_path = None
        profile.updated_at = None

        status = parser.get_resume_status(profile)

        assert status["has_resume"] is False
        assert status["has_file"] is False

    def test_with_resume_status(self):
        """Test status when resume is uploaded."""
        from app.services.resume_parser import ResumeParserService

        parser = ResumeParserService()
        profile = MagicMock()
        profile.resume_text = "Resume content here" * 100
        profile.resume_file_path = "/resumes/test.pdf"
        profile.updated_at = MagicMock()
        profile.updated_at.isoformat.return_value = "2026-07-05T12:00:00"

        status = parser.get_resume_status(profile)

        assert status["has_resume"] is True
        assert status["has_file"] is True
        assert status["resume_text_length"] > 0


# ─── Skill Prioritization Tests ────────────────────────────────────────────


class TestSkillPrioritization:
    """Test AEM skill prioritization in the list."""

    def test_aem_skills_come_first(self):
        """Test that AEM-related skills are sorted to the front."""
        from app.services.resume_parser import ResumeParserService

        parser = ResumeParserService()
        skills = ["Docker", "AEM", "Python", "OSGi", "Jenkins", "Sling", "React"]
        result = parser._prioritize_skills(skills)

        # AEM-related skills should be first
        aem_keywords = {"aem", "sling", "osgi"}
        aem_skills = [s for s in result if any(kw in s.lower() for kw in aem_keywords)]
        other_skills = [s for s in result if not any(kw in s.lower() for kw in aem_keywords)]

        # All AEM skills should come before non-AEM skills
        last_aem_idx = max(result.index(s) for s in aem_skills)
        if other_skills:
            first_other_idx = min(result.index(s) for s in other_skills)
            assert last_aem_idx < first_other_idx
