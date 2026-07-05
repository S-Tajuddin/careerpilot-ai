"""
CareerPilot AI — Resume Parser Service
Extracts text from PDF/DOCX files, then uses Gemini Flash
to parse skills, experience, summary, and job preferences.

Flow:
  1. Upload file (PDF/DOCX/TXT) → save to data/resumes/
  2. Extract raw text from file
  3. Send to Gemini Flash for structured extraction
  4. Update profile with parsed data (skills, experience, summary, etc.)
  5. Re-score all existing jobs against the updated profile
"""

import json
import os
import re
import shutil
from datetime import datetime, timezone
from pathlib import Path
from typing import Optional

from sqlalchemy.orm import Session

from app.config import settings, RESUMES_DIR
from app.models import Profile, Job
from app.services.llm import llm_service


# Common AEM/tech keywords for heuristic skill extraction when LLM output is truncated
_SKILL_KEYWORDS = [
    "AEM", "AEM 6.5", "AEM as a Cloud Service", "AEM Cloud Service", "Adobe Experience Manager",
    "Edge Delivery Services", "EDS", "Universal Editor", "Sling", "OSGi", "HTL", "Sightly",
    "Dispatcher", "CQ5", "Java", "JavaScript", "TypeScript", "React", "Angular", "Node.js",
    "Maven", "Gradle", "Git", "Jenkins", "Docker", "Kubernetes", "AWS", "Azure",
    "JCR", "CRX", "Content Fragments", "Experience Fragments", "Core Components",
    "GraphQL", "REST", "SOAP", "Spring", "JUnit", "Mockito", "SQL", "MySQL", "MongoDB",
    "HTML", "CSS", "SCSS", "Tailwind", "Webpack", "npm", "CI/CD", "Agile", "Scrum",
    "Franklin", "Adobe Analytics", "Adobe Target", "Launch", "Solr", "Elasticsearch",
]


class ResumeParserService:
    """
    Parses uploaded resumes and updates the user profile.
    Uses Gemini Flash for intelligent extraction (free, fast, 500 RPD).
    """

    def __init__(self):
        self.resumes_dir = RESUMES_DIR
        self.resumes_dir.mkdir(parents=True, exist_ok=True)

    # ══════════════════════════════════════════════════════════════════════
    #  FILE TEXT EXTRACTION
    # ══════════════════════════════════════════════════════════════════════

    def extract_text_from_file(self, file_path: str) -> str:
        """Extract raw text from PDF, DOCX, or TXT file."""
        ext = Path(file_path).suffix.lower()

        if ext == ".pdf":
            return self._extract_pdf(file_path)
        elif ext in (".docx", ".doc"):
            return self._extract_docx(file_path)
        elif ext == ".txt":
            return self._extract_txt(file_path)
        else:
            raise ValueError(f"Unsupported file format: {ext}. Use PDF, DOCX, or TXT.")

    def _extract_pdf(self, file_path: str) -> str:
        """Extract text from PDF using PyPDF2."""
        try:
            from PyPDF2 import PdfReader
            reader = PdfReader(file_path)
            text_parts = []
            for page in reader.pages:
                page_text = page.extract_text()
                if page_text:
                    text_parts.append(page_text)
            return "\n\n".join(text_parts).strip()
        except Exception as e:
            raise ValueError(f"Failed to parse PDF: {e}")

    def _extract_docx(self, file_path: str) -> str:
        """Extract text from DOCX using python-docx."""
        try:
            from docx import Document
            doc = Document(file_path)
            text_parts = []
            for para in doc.paragraphs:
                if para.text.strip():
                    text_parts.append(para.text.strip())
            # Also extract from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            text_parts.append(cell.text.strip())
            return "\n".join(text_parts).strip()
        except Exception as e:
            raise ValueError(f"Failed to parse DOCX: {e}")

    def _extract_txt(self, file_path: str) -> str:
        """Read plain text file."""
        try:
            with open(file_path, "r", encoding="utf-8", errors="replace") as f:
                return f.read().strip()
        except Exception as e:
            raise ValueError(f"Failed to read text file: {e}")

    # ══════════════════════════════════════════════════════════════════════
    #  GEMINI-BASED STRUCTURED EXTRACTION
    # ══════════════════════════════════════════════════════════════════════

    async def parse_resume_with_llm(self, resume_text: str) -> dict:
        """
        Use Gemini Flash to extract structured data from resume text.
        Returns a dict with: full_name, email, phone, summary, skills,
        experience_years, current_role, current_company, education,
        certifications, target_roles, preferred_work_arrangement
        """
        prompt = f"""You are an expert resume parser for an AEM/EDS Developer career tool. 
Parse this resume and extract structured information.

RESUME TEXT:
---
{resume_text[:8000]}
---

Extract the following and return as a JSON object with these EXACT keys:
{{
  "full_name": "string or null",
  "email": "string or null", 
  "phone": "string or null",
  "summary": "A concise 2-3 sentence professional summary based on the resume. Highlight AEM/EDS expertise.",
  "skills": ["List of ALL technical skills mentioned - include specific versions, tools, frameworks. Be comprehensive."],
  "experience_years": number (total years of professional experience, decimal ok),
  "current_role": "Current or most recent job title",
  "current_company": "Current or most recent employer name",
  "education": "Highest degree and institution, e.g. 'B.Tech CSE, JNTU'",
  "certifications": ["List of certifications if any"],
  "target_roles": ["What roles this person should target based on their experience - e.g. 'Senior AEM Developer', 'AEM Architect'"],
  "preferred_work_arrangement": "remote/hybrid/onsite/any — based on resume clues or default 'any'",
  "notable_projects": ["Brief descriptions of key AEM/EDS projects mentioned"],
  "location": "Current location if mentioned, otherwise null",
  "companies_worked": ["List of companies worked at"]
}}

IMPORTANT RULES:
- For skills, be VERY thorough — extract every technology, framework, tool, platform mentioned
- Include both primary skills (AEM, Java, etc.) and secondary skills (Git, Jenkins, etc.)
- If a skill version is mentioned (e.g. "AEM 6.5"), include both "AEM" and "AEM 6.5"
- For experience_years, calculate from work history dates if available
- For target_roles, suggest based on experience level and skills
- If something is not found in the resume, use null (not a made-up value)
- DO NOT fabricate or infer skills not mentioned in the resume

Output ONLY the JSON object, no markdown fences, no commentary."""

        response, model_used = await llm_service.generate(
            prompt=prompt,
            system="You are a precise resume parser. Always respond with valid JSON only. No markdown, no explanations.",
            task_type="interactive",  # Use Gemini Flash for this
            temperature=0.1,
            max_tokens=8192,
        )

        parsed = self._parse_llm_response(response)
        heuristics = self.extract_heuristics_from_text(resume_text)
        return self._merge_parsed_with_heuristics(parsed, heuristics)

    def extract_heuristics_from_text(self, resume_text: str) -> dict:
        """Extract contact info and skills from raw resume text when LLM parsing fails or is incomplete."""
        text = resume_text or ""
        normalized = re.sub(r"\s+", " ", text)
        result: dict = {}

        email_match = re.search(r"[\w.+-]+@[\w.-]+\.[A-Za-z]{2,}", text)
        if email_match:
            result["email"] = email_match.group(0).strip()

        phone_match = re.search(
            r"(?:\+?\d{1,3}[\s-]?)?(?:\(?\d{2,4}\)?[\s-]?)?\d{3,4}[\s-]?\d{4,6}",
            text,
        )
        if phone_match:
            phone = re.sub(r"\s+", " ", phone_match.group(0)).strip()
            digits = re.sub(r"\D", "", phone)
            if len(digits) >= 10:
                result["phone"] = phone

        for line in text.splitlines()[:8]:
            candidate = re.sub(r"\s+", " ", line).strip()
            if not candidate or len(candidate) > 80:
                continue
            lower = candidate.lower()
            if any(x in lower for x in ("@", "http", "linkedin", "summary", "resume", "curriculum")):
                continue
            if re.search(r"\d{5,}", candidate):
                continue
            if re.match(r"^[A-Za-z][A-Za-z\s.'-]{2,60}$", candidate) and 2 <= len(candidate.split()) <= 5:
                result["full_name"] = candidate.title() if candidate.isupper() else candidate
                break

        exp_match = re.search(
            r"(?:around|over|more than|approximately|~)?\s*(\d+(?:\.\d+)?)\+?\s*(?:years?|yrs?)\s+(?:of\s+)?(?:experience|exp)",
            normalized,
            re.IGNORECASE,
        )
        if exp_match:
            result["experience_years"] = float(exp_match.group(1))

        role_match = re.search(
            r"(Senior\s+)?(?:AEM|Adobe Experience Manager|EDS|Edge Delivery)[\w\s/\-]{0,40}(?:Developer|Engineer|Architect|Consultant|Lead)",
            text,
            re.IGNORECASE,
        )
        if role_match:
            result["current_role"] = re.sub(r"\s+", " ", role_match.group(0)).strip()

        location_match = re.search(
            r"(?:location|based in|city)[:\s]+([A-Za-z\s,]+(?:India|IN))",
            text,
            re.IGNORECASE,
        )
        if location_match:
            result["location"] = location_match.group(1).strip()
        elif re.search(r"\bHyderabad\b", text, re.IGNORECASE):
            result["location"] = "Hyderabad, India"

        skills = []
        text_lower = text.lower()
        for skill in _SKILL_KEYWORDS:
            if skill.lower() in text_lower and skill not in skills:
                skills.append(skill)
        if skills:
            result["skills"] = skills

        if "architect" in text_lower and "aem" in text_lower:
            result["target_roles"] = ["AEM Architect", "Senior AEM Developer"]
        elif "senior" in text_lower and "aem" in text_lower:
            result["target_roles"] = ["Senior AEM Developer", "AEM Architect"]

        return result

    def _merge_parsed_with_heuristics(self, parsed: dict, heuristics: dict) -> dict:
        """Fill missing LLM fields with heuristic extractions."""
        merged = dict(heuristics)
        merged.update({k: v for k, v in (parsed or {}).items() if v not in (None, "", [], {})})
        for key in ("full_name", "email", "phone", "summary", "current_role", "current_company", "location"):
            if not merged.get(key) and heuristics.get(key):
                merged[key] = heuristics[key]
        if not merged.get("skills") and heuristics.get("skills"):
            merged["skills"] = heuristics["skills"]
        elif merged.get("skills") and heuristics.get("skills"):
            merged["skills"] = list(dict.fromkeys(merged["skills"] + heuristics["skills"]))
        if merged.get("experience_years") is None and heuristics.get("experience_years") is not None:
            merged["experience_years"] = heuristics["experience_years"]
        if not merged.get("target_roles") and heuristics.get("target_roles"):
            merged["target_roles"] = heuristics["target_roles"]
        return merged

    def _parse_llm_response(self, response: str) -> dict:
        """Parse LLM response into structured dict."""
        if not response:
            return {}

        # Try direct JSON parse
        try:
            result = json.loads(response)
            if isinstance(result, dict):
                return result
        except json.JSONDecodeError:
            pass

        # Try extracting JSON from markdown fences
        import re
        json_match = re.search(r'```(?:json)?\s*([\s\S]*?)```', response)
        if json_match:
            try:
                result = json.loads(json_match.group(1).strip())
                if isinstance(result, dict):
                    return result
            except json.JSONDecodeError:
                pass

        # Try finding JSON object in text
        brace_match = re.search(r'\{[\s\S]*\}', response)
        if brace_match:
            try:
                result = json.loads(brace_match.group())
                if isinstance(result, dict):
                    return result
            except json.JSONDecodeError:
                pass

        # Recover individual fields from truncated JSON
        recovered = self._recover_partial_json(response)
        if recovered:
            return recovered

        print(f"Failed to parse LLM resume response: {response[:200]}...")
        return {}

    def _recover_partial_json(self, response: str) -> dict:
        """Extract fields from truncated or malformed JSON LLM output."""
        if not response or "{" not in response:
            return {}

        result: dict = {}
        string_fields = (
            "full_name", "email", "phone", "summary", "current_role",
            "current_company", "education", "location", "preferred_work_arrangement",
        )
        for field in string_fields:
            match = re.search(
                rf'"{field}"\s*:\s*"((?:[^"\\]|\\.)*)"',
                response,
                re.DOTALL,
            )
            if match:
                value = match.group(1).replace('\\"', '"').replace("\\n", "\n").strip()
                if value and value.lower() != "null":
                    result[field] = value

        exp_match = re.search(r'"experience_years"\s*:\s*(\d+(?:\.\d+)?)', response)
        if exp_match:
            result["experience_years"] = float(exp_match.group(1))

        for list_field in ("skills", "certifications", "target_roles", "notable_projects", "companies_worked"):
            match = re.search(rf'"{list_field}"\s*:\s*\[(.*?)\]', response, re.DOTALL)
            if not match:
                continue
            items = re.findall(r'"((?:[^"\\]|\\.)*)"', match.group(1))
            if items:
                result[list_field] = [i.replace('\\"', '"') for i in items]

        return result

    # ══════════════════════════════════════════════════════════════════════
    #  PROFILE UPDATE FROM PARSED RESUME
    # ══════════════════════════════════════════════════════════════════════

    def update_profile_from_parsed_resume(
        self, parsed: dict, resume_text: str, file_path: str, profile: Profile, db: Session
    ) -> dict:
        """
        Update the Profile model with parsed resume data.
        Resume is the SOURCE OF TRUTH — always overwrites with parsed values.
        Only skips fields where the parsed value is null/empty.
        Returns a summary of what was updated.
        """
        updated_fields = []

        # ── Always overwrite from resume (resume = source of truth) ──────

        if parsed.get("full_name"):
            profile.full_name = parsed["full_name"]
            updated_fields.append("full_name")

        if parsed.get("email"):
            profile.email = parsed["email"]
            updated_fields.append("email")

        if parsed.get("phone"):
            profile.phone = parsed["phone"]
            updated_fields.append("phone")

        # Summary — always update from resume
        if parsed.get("summary"):
            profile.summary = parsed["summary"]
            updated_fields.append("summary")

        # Skills — MERGE resume skills with existing profile skills
        resume_skills = parsed.get("skills", [])
        if resume_skills:
            existing_skills = profile.skills_list if hasattr(profile, 'skills_list') else []
            # Merge: add new skills from resume, keep existing ones not in resume
            merged = list(set(existing_skills + resume_skills))
            # Sort: AEM/EDS related first
            profile.skills_list = self._prioritize_skills(merged)
            updated_fields.append("skills")

        # Experience — always update from resume
        if parsed.get("experience_years") is not None:
            profile.experience_years = float(parsed["experience_years"])
            updated_fields.append("experience_years")

        # Current role — always update from resume
        if parsed.get("current_role"):
            profile.current_role = parsed["current_role"]
            updated_fields.append("current_role")

        # Target role — always update from target_roles list
        if parsed.get("target_roles") and len(parsed["target_roles"]) > 0:
            profile.target_role = parsed["target_roles"][0]
            updated_fields.append("target_role")

        # Location — always update from resume
        if parsed.get("location"):
            profile.location = parsed["location"]
            updated_fields.append("location")

        # Work arrangement — always update from resume
        if parsed.get("preferred_work_arrangement"):
            profile.remote_preference = parsed["preferred_work_arrangement"]
            updated_fields.append("remote_preference")

        # Store resume text and file path
        profile.resume_text = resume_text
        profile.resume_file_path = file_path
        updated_fields.extend(["resume_text", "resume_file_path"])

        # Store extra parsed data as JSON appended to summary
        extra_data = {}
        if parsed.get("education"):
            extra_data["education"] = parsed["education"]
        if parsed.get("current_company"):
            extra_data["current_company"] = parsed["current_company"]
        if parsed.get("certifications"):
            extra_data["certifications"] = parsed["certifications"]
        if parsed.get("notable_projects"):
            extra_data["notable_projects"] = parsed["notable_projects"]
        if parsed.get("companies_worked"):
            extra_data["companies_worked"] = parsed["companies_worked"]
        if parsed.get("target_roles"):
            extra_data["target_roles"] = parsed["target_roles"]

        # Replace any previous [PARSED_DATA] block, then append new one
        if extra_data:
            import re
            # Strip old PARSED_DATA block if present
            if profile.summary:
                profile.summary = re.sub(r'\n*\[PARSED_DATA\].*?\[/PARSED_DATA\]', '', profile.summary, flags=re.DOTALL).strip()
            extra_json = json.dumps(extra_data, ensure_ascii=False)
            profile.summary = f"{profile.summary}\n\n[PARSED_DATA]{extra_json}[/PARSED_DATA]" if profile.summary else f"[PARSED_DATA]{extra_json}[/PARSED_DATA]"

        profile.updated_at = datetime.now(timezone.utc)
        db.commit()
        db.refresh(profile)

        return {
            "updated_fields": updated_fields,
            "skills_count": len(profile.skills_list) if hasattr(profile, 'skills_list') else 0,
            "experience_years": profile.experience_years,
            "current_role": profile.current_role,
            "target_role": profile.target_role,
            "full_name": profile.full_name,
            "email": profile.email,
            "phone": profile.phone,
            "location": profile.location,
            "remote_preference": profile.remote_preference,
        }

    def _prioritize_skills(self, skills: list[str]) -> list[str]:
        """Sort skills: AEM/EDS first, then others."""
        aem_keywords = {"aem", "eds", "adobe experience manager", "edge delivery",
                        "sling", "osgi", "htl", "sightly", "cq5", "dispatcher"}
        priority = []
        others = []
        for skill in skills:
            if any(kw in skill.lower() for kw in aem_keywords):
                priority.append(skill)
            else:
                others.append(skill)
        return priority + others

    # ══════════════════════════════════════════════════════════════════════
    #  FULL PIPELINE: Upload → Parse → Update Profile → Re-score
    # ══════════════════════════════════════════════════════════════════════

    async def process_resume_upload(
        self, file_content: bytes, filename: str, profile: Profile, db: Session
    ) -> dict:
        """
        Full pipeline: save file → extract text → parse with LLM → update profile.
        Returns summary of what was updated.
        """
        # 1. Save file
        timestamp = datetime.now(timezone.utc).strftime("%Y%m%d_%H%M%S")
        safe_name = Path(filename).stem.replace(" ", "_")
        ext = Path(filename).suffix.lower()
        saved_filename = f"{safe_name}_{timestamp}{ext}"
        saved_path = self.resumes_dir / saved_filename

        with open(saved_path, "wb") as f:
            f.write(file_content)

        # 2. Extract text
        try:
            resume_text = self.extract_text_from_file(str(saved_path))
        except Exception as e:
            # Clean up file on parse error
            saved_path.unlink(missing_ok=True)
            raise ValueError(f"Could not extract text from {filename}: {e}")

        if not resume_text or len(resume_text.strip()) < 50:
            saved_path.unlink(missing_ok=True)
            raise ValueError(f"Resume appears empty or too short ({len(resume_text)} chars). Please upload a valid resume.")

        # 3. Parse with LLM (+ heuristic fallbacks)
        parsed = await self.parse_resume_with_llm(resume_text)

        if not parsed:
            parsed = self.extract_heuristics_from_text(resume_text)

        if not parsed:
            # Still save the text even if LLM parsing failed
            profile.resume_text = resume_text
            profile.resume_file_path = str(saved_path)
            db.commit()
            return {
                "updated_fields": ["resume_text", "resume_file_path"],
                "warning": "LLM parsing failed — resume text saved but structured extraction incomplete. Try re-uploading later.",
            }

        # 4. Update profile
        update_result = self.update_profile_from_parsed_resume(
            parsed=parsed,
            resume_text=resume_text,
            file_path=str(saved_path),
            profile=profile,
            db=db,
        )

        # 5. Re-score all existing jobs with the updated profile
        await self._rescore_all_jobs(profile, db)

        return {
            **update_result,
            "parsed_data": parsed,
            "resume_text_length": len(resume_text),
            "file_saved": str(saved_path),
        }

    async def reparse_stored_resume(self, profile: Profile, db: Session) -> dict:
        """Re-parse an already-uploaded resume and refresh profile fields."""
        if not profile.resume_text or len(profile.resume_text.strip()) < 50:
            raise ValueError("No resume text stored — upload a resume first.")

        parsed = await self.parse_resume_with_llm(profile.resume_text)
        if not parsed:
            parsed = self.extract_heuristics_from_text(profile.resume_text)

        file_path = profile.resume_file_path or ""
        update_result = self.update_profile_from_parsed_resume(
            parsed=parsed,
            resume_text=profile.resume_text,
            file_path=file_path,
            profile=profile,
            db=db,
        )
        await self._rescore_all_jobs(profile, db)

        return {
            **update_result,
            "parsed_data": parsed,
            "message": "Profile updated from stored resume",
        }

    async def _rescore_all_jobs(self, profile: Profile, db: Session):
        """Quick re-score all active jobs with the updated profile."""
        try:
            from app.services.scoring import scoring_service

            active_jobs = db.query(Job).filter(Job.is_active == True).all()
            scored = 0
            for job in active_jobs:
                try:
                    result = scoring_service.quick_score(job, profile, db)
                    job.match_score = result["overall_score"]
                    job.match_strengths_list = result["strengths"]
                    job.match_weaknesses_list = result["weaknesses"]
                    job.match_recommendations_list = result["recommendations"]
                    scored += 1
                except Exception as e:
                    print(f"Re-score error for job {job.id}: {e}")

            db.commit()
            if scored:
                print(f"  ⚡ Re-scored {scored} jobs with updated resume profile")
        except Exception as e:
            print(f"Re-scoring failed: {e}")

    # ══════════════════════════════════════════════════════════════════════
    #  RESUME-BASED SEARCH QUERY GENERATION
    # ══════════════════════════════════════════════════════════════════════

    def generate_search_queries_from_resume(self, profile: Profile) -> list[dict]:
        """
        Generate smart search queries based on resume content.
        Returns list of {query, location} dicts for auto-searching.
        """
        queries = []

        skills = profile.skills_list if hasattr(profile, 'skills_list') else []
        target_role = profile.target_role or "AEM Developer"
        current_role = profile.current_role or "AEM Developer"
        location = profile.location or "Hyderabad, India"

        # Core role-based searches
        queries.append({"query": f"{target_role}", "location": "India", "label": f"{target_role}"})
        queries.append({"query": f"{target_role} remote", "location": "", "label": f"{target_role} Remote"})

        # AEM-specific searches from skills
        aem_skills = [s for s in skills if any(kw in s.lower() for kw in [
            "aem", "adobe experience", "cq5", "eds", "edge delivery",
            "sling", "osgi", "htl", "sightly", "dispatcher"
        ])]

        if aem_skills:
            top_aem = aem_skills[:3]
            queries.append({
                "query": f"{' '.join(top_aem)} developer",
                "location": "India",
                "label": f"{' + '.join(top_aem[:2])} Dev"
            })

        # Title variations
        title_variations = self._get_title_variations(target_role)
        for variant in title_variations:
            queries.append({"query": variant, "location": "India", "label": variant})

        # Location-specific
        city = location.split(",")[0].strip() if location else "Hyderabad"
        queries.append({"query": f"AEM developer {city}", "location": city, "label": f"AEM in {city}"})

        # Senior/Architect level
        exp = profile.experience_years or 8
        if exp >= 8:
            queries.append({"query": "senior AEM architect", "location": "India", "label": "Sr AEM Architect"})
            queries.append({"query": "AEM tech lead", "location": "India", "label": "AEM Tech Lead"})
        if exp >= 10:
            queries.append({"query": "AEM principal engineer", "location": "India", "label": "AEM Principal Eng"})
            queries.append({"query": "Adobe Experience Manager architect remote", "location": "", "label": "AEM Architect Remote"})

        # EDS-specific if in skills
        if any("eds" in s.lower() or "edge delivery" in s.lower() for s in skills):
            queries.append({"query": "Edge Delivery Services developer", "location": "India", "label": "EDS Developer"})
            queries.append({"query": "AEM Edge Delivery engineer", "location": "India", "label": "AEM EDS Engineer"})

        # Remove duplicates while preserving order
        seen = set()
        unique = []
        for q in queries:
            key = q["query"].lower()
            if key not in seen:
                seen.add(key)
                unique.append(q)

        return unique[:8]  # Max 8 queries to avoid API rate limits

    def _get_title_variations(self, target_role: str) -> list[str]:
        """Generate title variations for search."""
        variations = []
        role_lower = target_role.lower()

        if "aem" in role_lower:
            variations.extend([
                "AEM developer",
                "Adobe Experience Manager developer",
                "CQ5 developer",
                "AEM Sites developer",
            ])
        if "architect" in role_lower:
            variations.append("AEM solution architect")
        if "eds" in role_lower or "edge" in role_lower:
            variations.extend([
                "EDS developer",
                "AEM Franklin developer",
            ])

        return variations[:4]

    # ══════════════════════════════════════════════════════════════════════
    #  GET RESUME STATUS
    # ══════════════════════════════════════════════════════════════════════

    def get_resume_status(self, profile: Profile) -> dict:
        """Get resume upload status and summary."""
        has_resume = bool(profile.resume_text)
        has_file = bool(profile.resume_file_path)

        return {
            "has_resume": has_resume,
            "has_file": has_file,
            "file_path": profile.resume_file_path if has_file else None,
            "resume_text_length": len(profile.resume_text) if profile.resume_text else 0,
            "last_updated": profile.updated_at.isoformat() if profile.updated_at else None,
            "skills_from_resume": has_resume and len(profile.skills_list or []) > 0,
        }


# Singleton
resume_parser = ResumeParserService()
